import os
from docx import Document

def write_folder_contents_to_word(folder_path, output_file):
    document = Document()
    
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            file_path = os.path.join(root, file)
            document.add_heading(file_path, level=1)
            with open(file_path, 'r', encoding='utf-8') as f:
                if '__pycache__' not in root:
                    content = f.read()
                document.add_paragraph(content)
    
    document.save(output_file)

folder_path = '/Users/treasures_y/Documents/code/HG/AutoML/python/automl/alserver/codecs'
output_file = '/Users/treasures_y/Documents/code/HG/AutoML/file.docx'
write_folder_contents_to_word(folder_path, output_file)